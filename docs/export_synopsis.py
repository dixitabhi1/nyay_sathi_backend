from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "NyayaSetu_Synopsis.md"
DOCX_OUTPUT = ROOT / "NyayaSetu_Synopsis.docx"
PDF_OUTPUT = ROOT / "NyayaSetu_Synopsis.pdf"


def _clean_inline(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`(.*?)`", r"\1", text)
    return text.strip()


def parse_markdown(path: Path) -> list[tuple[str, str]]:
    lines = path.read_text(encoding="utf-8").splitlines()
    blocks: list[tuple[str, str]] = []
    paragraph: list[str] = []
    bullets: list[str] = []

    def flush_paragraph() -> None:
        if paragraph:
            blocks.append(("paragraph", " ".join(part.strip() for part in paragraph if part.strip())))
            paragraph.clear()

    def flush_bullets() -> None:
        if bullets:
            for bullet in bullets:
                blocks.append(("bullet", bullet))
            bullets.clear()

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            flush_bullets()
            continue
        if stripped.startswith("# "):
            flush_paragraph()
            flush_bullets()
            blocks.append(("title", _clean_inline(stripped[2:])))
            continue
        if stripped.startswith("## "):
            flush_paragraph()
            flush_bullets()
            blocks.append(("heading2", _clean_inline(stripped[3:])))
            continue
        if stripped.startswith("### "):
            flush_paragraph()
            flush_bullets()
            blocks.append(("heading3", _clean_inline(stripped[4:])))
            continue
        if stripped.startswith("- "):
            flush_paragraph()
            bullets.append(_clean_inline(stripped[2:]))
            continue
        paragraph.append(_clean_inline(stripped))

    flush_paragraph()
    flush_bullets()
    return blocks


def export_docx(blocks: list[tuple[str, str]], output_path: Path) -> None:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for block_type, text in blocks:
        if block_type == "title":
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(18)
        elif block_type == "heading2":
            document.add_heading(text, level=1)
        elif block_type == "heading3":
            document.add_heading(text, level=2)
        elif block_type == "bullet":
            document.add_paragraph(text, style="List Bullet")
        else:
            document.add_paragraph(text)

    document.save(output_path)


def export_pdf(blocks: list[tuple[str, str]], output_path: Path) -> None:
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "SynopsisTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        alignment=1,
        spaceAfter=14,
    )
    h1_style = ParagraphStyle(
        "SynopsisH1",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=18,
        spaceBefore=8,
        spaceAfter=8,
    )
    h2_style = ParagraphStyle(
        "SynopsisH2",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=16,
        spaceBefore=6,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "SynopsisBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=14,
        spaceAfter=8,
    )

    story = []
    pending_bullets: list[ListItem] = []

    def flush_bullets() -> None:
        nonlocal pending_bullets
        if pending_bullets:
            story.append(
                ListFlowable(
                    pending_bullets,
                    bulletType="bullet",
                    start="circle",
                    leftIndent=16,
                )
            )
            story.append(Spacer(1, 8))
            pending_bullets = []

    for block_type, text in blocks:
        if block_type == "bullet":
            pending_bullets.append(ListItem(Paragraph(text, body_style)))
            continue

        flush_bullets()
        if block_type == "title":
            story.append(Paragraph(text, title_style))
        elif block_type == "heading2":
            story.append(Paragraph(text, h1_style))
        elif block_type == "heading3":
            story.append(Paragraph(text, h2_style))
        else:
            story.append(Paragraph(text, body_style))

    flush_bullets()
    doc.build(story)


def main() -> None:
    blocks = parse_markdown(SOURCE)
    export_docx(blocks, DOCX_OUTPUT)
    export_pdf(blocks, PDF_OUTPUT)
    print(f"Wrote {DOCX_OUTPUT}")
    print(f"Wrote {PDF_OUTPUT}")


if __name__ == "__main__":
    main()
